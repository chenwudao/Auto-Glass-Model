from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt # 用于设置字体大小
from docx.oxml.ns import qn # 用于设置中文字体

def generate_project_summary_doc(summary_content, filename="项目工作总结.docx"):
    """
    根据提供的总结内容生成一个Word文档。

    Args:
        summary_content (dict): 包含总结各部分内容的字典。
        filename (str): 生成的Word文档的文件名。
    """
    document = Document()

    # 设置默认字体，以支持中文显示
    # 对于Windows系统，通常使用'宋体'或'微软雅黑'
    # 对于macOS，可能需要设置'PingFang SC'或'Heiti SC'
    # 如果遇到乱码，请尝试更换字体或查找系统支持的中文字体名称
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    document.add_heading(summary_content.get("title", "个人工作总结"), level=0)
    document.add_paragraph() # 添加一个空行

    # 添加项目名称和岗位
    document.add_heading('项目名称：' + summary_content.get("project_name", "N/A"), level=2)
    document.add_heading('岗位：' + summary_content.get("position", "N/A"), level=2)
    document.add_heading('总结周期：' + summary_content.get("period", "N/A"), level=2)
    document.add_paragraph()

    # 一、本月/项目阶段承担的工作任务及交付情况
    document.add_heading('一、本月/项目阶段承担的工作任务及交付情况', level=1)
    document.add_paragraph('在本月/项目阶段，我作为星眸智能项目的核心算法技术人员，主要承担了以下工作任务：')
    for task in summary_content.get("tasks", []):
        document.add_paragraph(task, style='List Bullet') # 项目符号列表

    document.add_paragraph('**交付物是否按时提交：**' + summary_content.get("delivery_status", "N/A"), style='Normal')
    document.add_paragraph()

    # 二、产出成果与创新贡献
    document.add_heading('二、产出成果与创新贡献', level=1)
    document.add_paragraph('本阶段我主要的产出成果为：')
    for result in summary_content.get("results", []):
        document.add_paragraph(result, style='List Bullet')

    document.add_paragraph()

    # 三、会议参与及团队协作情况
    document.add_heading('三、会议参与及团队协作情况', level=1)
    for meeting_info in summary_content.get("meetings_collaboration", []):
        document.add_paragraph(meeting_info)

    document.add_paragraph()

    # 四、自省与思考
    document.add_heading('四、自省与思考', level=1)
    document.add_paragraph(summary_content.get("reflection", ""))

    # 保存文档
    try:
        document.save(filename)
        print(f"Word 文档 '{filename}' 已成功生成。")
    except Exception as e:
        print(f"生成文档时发生错误: {e}")

# 示例使用
if __name__ == "__main__":
    # 使用你之前提供的总结内容来填充这个字典
    my_summary_content = {
        "title": "星眸智能项目个人工作总结",
        "project_name": "互联网+星眸智能项目",
        "position": "算法技术人员",
        "period": "项目前期阶段（请在此处填写具体时间）", # 提醒用户填写具体时间
        "tasks": [
            "算法建模思路的提供与核心价值点的提炼：积极参与项目技术方案的讨论与构建，负责算法建模方向的探索，并从“技术价值关键内容”和“技术优势”等维度，提炼出项目在动态眼轴预测、影响因素解析、微表情非侵入检测等方面的核心技术创新点。",
            "具体包括动态眼轴预测模型构建思路、影响因素量化解析算法思路、微表情非侵入检测算法思路。",
            "核心关键价值思路点的提炼：清晰阐述了星眸智能项目在解决“传统检测滞后”和“近视干预盲区”等核心挑战方面的创新价值，并将技术优势转化为用户可感知的亮点（如居家验配、多模式防控、视力趋势可视化和AI个性化指导）。",
            "项目技术文档内容的提供：根据与Bob师兄的两次会谈内容及项目技术需求，按时提供了关于核心技术（如SVR、深度学习、MLP、随机森林、YOLO算法等多模型协同，多模态特征工程，模型可解释性SHAP等）以及技术优势的文档内容。",
            "协助团队成员进行PPT制作：积极协助其他团队成员，将复杂的技术概念转化为清晰易懂的PPT内容，确保技术方案在对外展示中的准确性和专业性。"
        ],
        "delivery_status": "所有负责的技术思路提炼和文档内容均按照与Bob师兄沟通的节点按时交付，无任何拖延。",
        "results": [
            "技术文件内容产出：提供了“技术价值关键内容”和“技术优势”文档的核心技术和创新点部分，为项目技术方案的成型奠定了基础。",
            "算法建模思路创新：在和Bob师兄的讨论中，针对眼轴动态预测和微表情非侵入检测等方面，提出了具有创新性的算法建模思路，尤其是在多因素整合和微表情生理指标应用方面，为项目的技术突破提供了方向。",
            "协助优化技术表达：通过协助PPT制作，确保了技术内容的准确传达和亮点突出，提升了项目对外展示的效果。"
        ],
        "meetings_collaboration": [
            "会议参加情况：积极参与了与Bob师兄的两次关键会谈，并针对项目算法实现路径和技术创新点进行了深入讨论和反馈。",
            "是否在会议提出实质性内容：在会谈中，我针对如何整合多模态数据（环境、行为、生理、微表情）、选择合适的AI算法（如SVM、随机森林、MLP、YOLO等）以及实现模型可解释性（SHAP）等方面提出了具体建议和思路，这些建议被采纳并体现在最终的技术方案中。",
            "参加的积极性与团队协作情况：在项目中保持了高度的积极性，与团队成员（特别是Bob师兄）沟通及时、响应迅速，确保了信息流通顺畅。在协作方面，我主动提供技术支持，协助其他成员完成PPT制作，展现了良好的团队合作精神。"
        ],
        "reflection": "回顾本阶段的工作，我在深入研究算法建模、实现大量代码、查阅前沿文献以及构建严谨数学模型方面投入了大量精力，并取得了一系列技术成果。我深知，作为算法技术人员，不仅要专注于技术本身，更要学会如何有效地将这些复杂的、具有核心价值的技术成果进行阐释和转化，使其能够被团队其他成员和项目决策者充分理解和利用。\n\n因此，在未来工作中，我将更加主动地参与到与我个人技术成果相关的PPT制作与汇报工作中。我相信，只有我最了解这些算法模型背后的逻辑和代码实现细节，也才能最准确、最精炼地向他人传达其核心价值和创新点。我将不断提升自身在技术表达和沟通方面的能力，努力将深奥的技术转化为易于理解和应用的成果，从而更好地赋能团队，为项目的成功贡献更多力量。"
    }

    generate_project_summary_doc(my_summary_content)